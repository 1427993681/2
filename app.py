import sys
import pandas as pd
from docx import Document

def main(excel_path):
    # 读取 Excel 数据
    df = pd.read_excel(excel_path)
    
    # TODO: 根据实际列名进行统计分析
    # 示例：统计各类案件数量
    summary = df.describe(include='all')
    
    # 生成 Word 报告
    doc = Document()
    doc.add_heading('全市犯罪情况统计报告', level=1)
    doc.add_paragraph(f'数据来源：{excel_path}')
    doc.add_paragraph('一、电信诈骗案件概况')
    doc.add_paragraph(str(summary))
    doc.save('crime_report.docx')
    print("报告已生成：crime_report.docx")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("请提供 Excel 文件路径，例如：python app.py data.xlsx")
    else:
        main(sys.argv[1])
